#!/usr/bin/env python3
"""Generate two hand-off Word docs: code-vs-guidance change record + suggested
edits, one for the Brief (v4.5) and one for the Build Plan (v5.4). Dev/admin
tool — output lands in docs/, not served to participants."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

ACCENT = RGBColor(0xB5, 0x55, 0x2F)
MUTED = RGBColor(0x6A, 0x6A, 0x64)
DATE = datetime.date.today().isoformat()


def base_doc(title, subtitle):
    d = Document()
    st = d.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)
    h = d.add_heading(title, level=0)
    for run in h.runs: run.font.color.rgb = ACCENT
    p = d.add_paragraph(); r = p.add_run(subtitle); r.italic = True; r.font.color.rgb = MUTED
    meta = d.add_paragraph()
    rm = meta.add_run(f"Generated {DATE} from the current code on branch claude/peaceful-ramanujan-7iakx2. "
                      "Each row: what the guidance currently says → what the code now does → the suggested edit. "
                      "Apply these to bring the guidance document in line with the implementation.")
    rm.font.size = Pt(9); rm.font.color.rgb = MUTED
    return d


def h2(d, t):
    p = d.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def para(d, t, bold=False, italic=False, color=None, size=10.5):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    r.font.size = Pt(size); return p


def bullet(d, t):
    p = d.add_paragraph(style='List Bullet'); p.add_run(t); return p


def change_table(d, rows):
    """rows: list of (section, says_now, code_does, suggested)."""
    t = d.add_table(rows=1, cols=4); t.style = 'Light Grid Accent 2'
    hdr = t.rows[0].cells
    for i, label in enumerate(['Section', 'Guidance says now', 'Code now does', 'Suggested edit']):
        hdr[i].text = ''
        rr = hdr[i].paragraphs[0].add_run(label); rr.bold = True; rr.font.size = Pt(9)
    for sec, says, does, sug in rows:
        c = t.add_row().cells
        for i, val in enumerate([sec, says, does, sug]):
            c[i].text = ''
            rr = c[i].paragraphs[0].add_run(val); rr.font.size = Pt(8.5)
            if i == 3: rr.font.color.rgb = ACCENT
    for col, w in zip(range(4), [0.9, 2.1, 2.1, 2.4]):
        for cell in t.columns[col].cells: cell.width = Inches(w)
    return t


# ---- shared content -------------------------------------------------------
CIP_NEW_ITEMS = (
    "A — Commitment anxiety (career indecision; higher = more indecision): "
    "(1) “I can't commit to a career because I don't know what my other options are.” "
    "(2) “I am concerned that my career goals might change.” "
    "(3) “It is difficult to decide on a career because I like so many different things.”   "
    "B — Career decision self-efficacy / confidence (higher = more confident; scored forward): "
    "(1) “I am confident that I will be able to find a career.” "
    "(2) “I am quite confident that I will be able to find a career in which I'll perform well.” "
    "(3) “I am confident that I can overcome obstacles in pursuing my career.”   "
    "Six-point agreement scale (1 = completely disagree … 6 = strongly agree); shown in mixed order; "
    "pre + post; never shown to the AI; both scored as forward raw means (no reverse-keying)."
)

IMPLEMENTED = [
    ("(c) Skippable ~20s “imagine your future self” page before the mediator block", "Implemented (survey.jsx, sv-imagine sequence)."),
    ("(d) Role-play opens inside a concrete, specific scene (time/place/sensory)", "Implemented (lib/prompt.js main prompt, “open the very first message INSIDE a concrete scene”)."),
    ("(e) De-AI'd interface copy; no “Your future self, today” label", "Implemented (no such label remains)."),
    ("(f) Phase-B career picker as a separate page/modal returning to chat", "Implemented (pb-lock overlay; “Back to chat” preserves the conversation)."),
    ("(g) Vary response length/style across turns in BOTH stage-C prompts", "Implemented (main + baseline prompts both instruct deliberate length/rhythm variation)."),
    ("(h) Shrink “ideas to ask”; fixed four suggestions, no horizontal scroll, in the user's voice", "Implemented (AskIdeas: 2×2 grid, four items; mobile compacted/2-line clamp)."),
    ("(i) Light one-exchange location check in the direct prompt", "Implemented (guide Step-4 + geographic-realism floor in both phase-C prompts)."),
    ("(j) Pictorial closeness/continuity items use simpler participant-facing wording", "Implemented (IOS/FSCS plain wording)."),
    ("(k) Target population broadened to any major/year/university; no university field", "Implemented (pre-survey demographics + prompts generalised)."),
    ("(m) Voluntary interview shows the team email (thy.le@student.uva.nl); no contact collected; only a no-PII interest flag", "Implemented (post-survey Yes/No interview flag + email shown)."),
]

UX_POLISH = [
    "Comfort/display DEFAULT changed from maximal (A+++ / Roomy / Wide) to a neutral premium baseline (A+ / Cozy / Normal); all larger options remain available in the comfort panel (key bumped to v4).",
    "Frontend is now PRECOMPILED (build/*.js via `npm run build`) with React vendored locally — no runtime Babel and no CDN dependency to start; index.html loads the bundle (was: in-browser Babel + unpkg CDN).",
    "Phase-B no longer auto-opens the lock-in over the recommendation cards — cards stay in the thread; the chooser opens on a card tap or the “Choose a career” prompt.",
    "Each between-phase pause now offers a guarded “Back” to the previous phase (confirm where progress is lost); the post-survey still offers no path back into the role-play (§7).",
    "Cross-platform/UX hardening: mobile safe-area insets, 16px form fields (no iOS zoom), theme-aware scrollbars, button press feedback, rec-card hover lift, vertical-centering of short flow screens, mobile Likert fix (all 7 points reachable), decluttered mobile chat header + composer footer.",
    "Repo: GitHub Actions CI (npm ci → test → build), .nvmrc, .editorconfig, LICENSE.md added; superseded docs moved to docs/archive/; build_sync_test guards build/ against source drift.",
]


# ============================ BRIEF DOC ====================================
def build_brief():
    d = base_doc("Brief v4.5 → code: change record & suggested edits",
                 "What the implementation now does that diverges from Project Status Brief v4.5, with proposed wording.")

    h2(d, "1. Headline divergence — the distal outcome (CIP-Short)")
    para(d, "The Brief (v4.3–v4.5) consolidates the distal outcome to a SINGLE CIP-Short “Lack of Readiness” "
            "subscale (5 items, reverse-scored, 6-point), with H4 and H6 both built on it. The supervisor's "
            "confirmed scale note (“CIP_outcome_measures”) supersedes this: the code now implements TWO "
            "3-item outcomes, scored forward. This re-introduces two distal outcomes in place of the merged one "
            "and therefore changes the hypothesis wording — the main thing needing your/supervisor sign-off.", )
    para(d, "New instrument (as implemented):", bold=True)
    para(d, CIP_NEW_ITEMS)

    h2(d, "2. Change record")
    change_table(d, [
        ("§2.5 Outcomes & Measures",
         "Single instrument — CIP-Short Lack of Readiness; reverse-keyed confidence items index both lower indecision and higher self-efficacy.",
         "Two indices: commitment anxiety (cip_ca_1..3) and confidence/CDSE (cip_cf_1..3), 6-pt, forward-scored, mixed order, pre+post.",
         "Replace the single-LR description with the two indices; state both are forward-scored raw means; keep the original 6-point scale."),
        ("§2.6 H4 (distal)",
         "“Main shows a larger pre→post reduction in lack of readiness than baseline” (single, merged outcome).",
         "Two outcomes measured; anxiety expected to fall, confidence to rise (main vs baseline).",
         "Split H4 into H4a (commitment anxiety decreases more in main) and H4b (confidence increases more in main); or one H4 naming both."),
        ("§2.6 H6 (Andrea)",
         "Reflective vs directive on the LR subscale (α=.82).",
         "Same two CIP indices apply to the reflective-vs-direct contrast.",
         "Re-word H6 to the two indices (anxiety + confidence) instead of LR."),
        ("§3.4 / Appendix F (item sourcing)",
         "Five LR items; item 5 verbatim, items 1–4 reconstructed; LR doubles as self-efficacy proxy.",
         "Six items across CC (anxiety) + LR (confidence) blocks per the supervisor doc; loadings Xu (2020).",
         "Replace the 5-item LR table with the 6-item table from the supervisor doc; cite Hacker et al. (2013) + Xu (2020); keep the ‘confirm wording vs CIP-65 before fielding’ note."),
        ("§2.5 Response-scale note",
         "CIP-Short LR on 6 points.",
         "Both CIP indices on the same 6-point agreement scale.",
         "Keep 6-point; just refer to ‘the two CIP indices’ rather than ‘LR’."),
        ("RQ / framing (§2.x)",
         "‘career decision self-efficacy and career indecision’ served by one LR proxy.",
         "Now measured by two distinct indices (confidence; commitment anxiety).",
         "Keep the two-construct framing — it now maps cleanly onto the two measured indices; drop the ‘single proxy’ phrasing."),
    ])

    h2(d, "3. Suggested replacement wording (drop-in)")
    para(d, "§2.5 (measures):", bold=True)
    para(d, "“The two career-level aims are measured by two short indices drawn from the CIP-Short: commitment "
            "anxiety (3 items; career indecision) and career decision self-efficacy / confidence (3 items). Both use "
            "the original 6-point agreement scale (1 = completely disagree … 6 = strongly agree), are shown in mixed "
            "order, are scored forward (raw means, no reverse-keying), and are measured pre + post in both conditions.”")
    para(d, "§2.6 (H4):", bold=True)
    para(d, "“H4a (distal — indecision): the main condition shows a larger pre→post DECREASE in commitment "
            "anxiety than baseline. H4b (distal — self-efficacy): the main condition shows a larger pre→post "
            "INCREASE in confidence than baseline.” (H6 mirrors this for reflective vs directive.)")

    h2(d, "4. Already-implemented items the Brief can mark as done")
    for item, status in IMPLEMENTED:
        bullet(d, f"{item} — {status}")

    h2(d, "5. Other implementation notes (FYI; mostly app-side)")
    for u in UX_POLISH:
        bullet(d, u)

    out = "docs/Brief_v4.5_to_code_change_record_and_suggestions.docx"
    d.save(out); return out


# ============================ BUILD PLAN DOC ===============================
def build_plan():
    d = base_doc("Build Plan v5.4 → code: change record & suggested edits",
                 "What the implementation now does that diverges from Artifact Build Plan v5.4, with proposed wording.")

    h2(d, "1. Headline divergence — the distal outcome (CIP-Short)")
    para(d, "The Build Plan v5.4 still describes the single CIP-Short Lack-of-Readiness subscale (5 items, "
            "reverse-scored) and marks the CIP subscale “pending supervisor confirmation.” The supervisor's "
            "confirmed scale note resolves that pending item with a DIFFERENT instrument: two 3-item forward-scored "
            "outcomes. The code implements the confirmed version everywhere the old cip_lr lived.")
    para(d, "New instrument (as implemented):", bold=True)
    para(d, CIP_NEW_ITEMS)

    h2(d, "2. Change record")
    change_table(d, [
        ("§2 / §10.1i–j (distal outcome)",
         "Single CIP-Short LR subscale, reverse-scored; H4/H5 merged.",
         "Two forward outcomes: cip_ca (anxiety) + cip_cf (confidence).",
         "Describe the two indices; drop the merge note; align H-numbering with the Brief."),
        ("§10.1 Pre-survey block",
         "One CIP-LR block (5 items).",
         "One CIP block of 6 items (mixed order), 6-point, forward.",
         "Replace block contents with the 6 items; keep the single paged block."),
        ("§10.2 Post-survey",
         "post-CIP-LR (5).",
         "post-CIP (6) with _post ids.",
         "Replace with the 6 post items."),
        ("§14b / analysis",
         "One reverse-scored distal mean; one alpha.",
         "Two forward means (cip_anxiety, cip_confidence); two alphas; admin CSV + live descriptives split.",
         "Update the analysis description to two outcomes; note forward scoring and two reliabilities."),
        ("Instrument appendix",
         "CIP-Short LR (Xu & Tracey, 2017b); item 5 verbatim, 1–4 reconstructed.",
         "6 items from CC + LR blocks; loadings Xu (2020); Hacker et al. (2013) wording.",
         "Swap the appendix entry; keep ‘confirm against CIP-65 before fielding.’"),
        ("§5 Architecture (‘no-build React’)",
         "React/Babel from CDN; .jsx transpiled in-browser at load.",
         "Precompiled build/*.js + vendored React; no runtime Babel/CDN; npm run build; build_sync_test in CI.",
         "Change ‘no-build’ to ‘precompiled (no bundler)’; note the build step + CI guard."),
        ("§16 Comfort defaults",
         "Default A+++ / Roomy / Wide (text size MAX).",
         "Default A+ / Cozy / Normal (neutral premium); all larger options still in the panel; key v4.",
         "Note the neutral default; the maxima remain one tap away."),
        ("§7 Screen flow (Phase-B / pauses)",
         "Lock-in revealed when directions appear; pauses are forward-only.",
         "Cards stay visible (lock opens on tap); pauses now have a guarded Back; post-survey still no Back into role-play.",
         "Note the card-first behaviour + per-phase Back (with the §7 no-return-after-post rule kept)."),
    ])

    h2(d, "3. Suggested replacement wording (drop-in)")
    para(d, "§2 / §10 (distal outcome):", bold=True)
    para(d, "“The career-level distal outcomes are two short CIP-Short indices measured pre + post in both "
            "conditions and never fed to the model: commitment anxiety (cip_ca_1..3; higher = more career indecision) "
            "and career decision self-efficacy / confidence (cip_cf_1..3; higher = more confident). Both on the "
            "original 6-point agreement scale, shown mixed, scored forward (raw means).”")
    para(d, "§5 (architecture):", bold=True)
    para(d, "“Participant browser (precompiled React: build/*.js, vendored React) — source is plain .jsx, "
            "compiled by `npm run build`; no runtime Babel and no CDN dependency. CI runs npm test (incl. a "
            "build-sync guard) and the build.”")

    h2(d, "4. [TO-IMPLEMENT] items to mark IMPLEMENTED")
    for item, status in IMPLEMENTED:
        bullet(d, f"{item} — {status}")
    bullet(d, "(n) post-survey interview model: implemented in survey.jsx (Yes/No interview flag, team email shown, "
              "no contact collected); responses stored in the postSurvey JSONB — no schema column needed. The "
              "thesis §4.2 storage description should still be updated to match.")

    h2(d, "5. Other implementation notes (UX / repo)")
    for u in UX_POLISH:
        bullet(d, u)

    out = "docs/Build_Plan_v5.4_to_code_change_record_and_suggestions.docx"
    d.save(out); return out


if __name__ == "__main__":
    os.makedirs("docs", exist_ok=True)
    print("wrote", build_brief())
    print("wrote", build_plan())
